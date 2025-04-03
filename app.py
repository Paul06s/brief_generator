from flask import Flask, request, jsonify, send_file
from docx import Document
import json
import os
import re
from werkzeug.utils import secure_filename
import pytesseract
from PIL import Image

app = Flask(__name__)

# ---------------- Config/Paths ----------------
DATA_FOLDER = "data"
OUTPUT_FOLDER = "output"
UPLOAD_FOLDER = "uploads"

# Chaque fichier .docx correspond à un format
TEMPLATE_2VOLETS = "template_depliant_2volets.docx"
TEMPLATE_3VOLETS = "template_depliant_3volets.docx"
TEMPLATE_5VOLETS = "template_depliant_5volets.docx"
TEMPLATE_6VOLETS = "template_depliant_6volets.docx"
TEMPLATE_BROCHURE_16 = "template_brochure_16pages.docx"
TEMPLATE_CATALOGUE_24 = "template_catalogue_24pages.docx"

os.makedirs(DATA_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ---------------- OCR + Matching (exemple simple) ----------------
def extract_product_names_from_image(image_path):
    """Extrait du texte depuis l'image et renvoie une liste de lignes non vides."""
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image, lang='fra')
        lines = text.split("\n")
        return [line.strip() for line in lines if line.strip()]
    except Exception as e:
        app.logger.error(f"Erreur OCR: {e}")
        return []

def match_names_to_json(product_names):
    """Exemple : Cherche dans tous les .json du dossier data/"""
    matched_parfums = []
    json_files = [f for f in os.listdir(DATA_FOLDER) if f.endswith(".json")]
    for name in product_names:
        regex = re.compile(re.escape(name), re.IGNORECASE)
        for jf in json_files:
            path = os.path.join(DATA_FOLDER, jf)
            try:
                with open(path, encoding="utf-8") as f:
                    data = json.load(f)
                for parfum in data.get("parfums", []):
                    if regex.search(parfum['nom']):
                        matched_parfums.append(parfum)
            except Exception as e:
                app.logger.error(f"Erreur lecture JSON {jf}: {e}")
    return matched_parfums

# ---------------- Choix du template ----------------
def choose_template(type_doc):
    if type_doc == "depliant_2volets":
        return TEMPLATE_2VOLETS
    elif type_doc == "depliant_3volets":
        return TEMPLATE_3VOLETS
    elif type_doc == "depliant_5volets":
        return TEMPLATE_5VOLETS
    elif type_doc == "depliant_6volets":
        return TEMPLATE_6VOLETS
    elif type_doc == "brochure_16pages":
        return TEMPLATE_BROCHURE_16
    elif type_doc == "catalogue_24pages":
        return TEMPLATE_CATALOGUE_24
    else:
        # fallback
        return TEMPLATE_5VOLETS

# ---------------- Génération du Word ----------------
def fill_template(template_path, periode, parfums):
    """
    Exemple minimaliste : on ouvre le template, on fait un simple search & replace
    pour la période, puis on pourrait lister les parfums, etc.
    """
    doc = Document(template_path)

    # Search & replace basique
    for p in doc.paragraphs:
        if "{{periode}}" in p.text:
            p.text = p.text.replace("{{periode}}", periode)

    # Exemple : si tu veux ajouter à la fin la liste des parfums détectés
    if parfums:
        doc.add_heading("Parfums détectés", level=2)
        for parfum in parfums:
            doc.add_paragraph(parfum['nom'], style='List Bullet')
            # ... eventuellement + description etc.

    output_path = os.path.join(OUTPUT_FOLDER, f"brief_{periode.replace(' ', '')}.docx")
    doc.save(output_path)
    return output_path

# ---------------- Routes Flask ----------------
@app.route("/")
def index():
    # L'interface HTML avec la zone de collage et le menu déroulant
    return '''
    <html>
    <head>
        <title>Brief Generator</title>
        <style>
            body { font-family: sans-serif; margin: 40px; }
            #pasteZone {
                width: 400px;
                height: 200px;
                border: 2px dashed #ccc;
                padding: 10px;
                text-align: center;
                background-color: #f9f9f9;
                margin-bottom: 20px;
                overflow: hidden;
            }
            #pasteZone img {
                max-width: 100%;
                max-height: 100%;
                display: block;
                margin: auto;
            }
        </style>
    </head>
    <body>
        <h1>Générateur de Brief</h1>

        <!-- Zone de collage d'image pour OCR -->
        <div id="pasteZone" contenteditable="true">
            <p>Cliquez ici, puis CTRL+V / CMD+V pour coller l'image du chemin de fer</p>
        </div>

        <!-- Champ texte pour la période (ex: Avril 2025) -->
        <div>
            <label for="periode">Période (ex: Avril 2025) :</label>
            <input type="text" id="periode" name="periode">
        </div>

        <!-- Menu déroulant pour le type de doc -->
        <div>
            <label for="typeDoc">Type de document :</label>
            <select id="typeDoc">
                <option value="depliant_2volets">Dépliant 2 volets</option>
                <option value="depliant_3volets">Dépliant 3 volets</option>
                <option value="depliant_5volets" selected>Dépliant 5 volets</option>
                <option value="depliant_6volets">Dépliant 6 volets</option>
                <option value="brochure_16pages">Brochure 16 pages</option>
                <option value="catalogue_24pages">Catalogue 24 pages</option>
            </select>
        </div>

        <br>
        <button onclick="generateBrief()">Générer le brief Word</button>

        <script>
            let detectedProducts = [];

            // Événement paste (copier-coller d'image)
            document.getElementById('pasteZone').addEventListener('paste', function(e) {
                let items = (e.clipboardData || e.originalEvent.clipboardData).items;
                for (let i = 0; i < items.length; i++) {
                    let item = items[i];
                    if (item.kind === 'file') {
                        let file = item.getAsFile();
                        // On remplace le contenu de #pasteZone par l'image
                        let pasteZone = document.getElementById('pasteZone');
                        pasteZone.innerHTML = "";
                        let img = document.createElement("img");
                        img.src = URL.createObjectURL(file);
                        pasteZone.appendChild(img);
                        uploadFile(file);
                    }
                }
            });

            function uploadFile(file) {
                let formData = new FormData();
                formData.append("file", file);
                fetch("/upload_chemin", { method: "POST", body: formData })
                    .then(response => response.json())
                    .then(data => {
                        if (data.error) {
                            alert("Erreur : " + data.error);
                        } else {
                            alert("Produits détectés : " + data.produits_detectes.join(', '));
                            detectedProducts = data.produits_detectes;
                        }
                    })
                    .catch(err => alert("Erreur lors de l'upload : " + err));
            }

            function generateBrief() {
                let periode = document.getElementById("periode").value;
                let typeDoc = document.getElementById("typeDoc").value;

                if (!periode) {
                    alert("Veuillez indiquer la période !");
                    return;
                }

                fetch("/generate_brief", {
                    method: "POST",
                    headers: { "Content-Type": "application/json" },
                    body: JSON.stringify({
                        periode: periode,
                        typeDoc: typeDoc,
                        parfums: detectedProducts
                    })
                })
                .then(response => response.blob())
                .then(blob => {
                    let url = window.URL.createObjectURL(blob);
                    let a = document.createElement("a");
                    a.href = url;
                    a.download = "brief_" + periode.replace(/\\s+/g, '') + ".docx";
                    document.body.appendChild(a);
                    a.click();
                    a.remove();
                })
                .catch(err => alert("Erreur lors de la génération du brief : " + err));
            }
        </script>
    </body>
    </html>
    '''

@app.route("/upload_chemin", methods=["POST"])
def upload_chemin():
    """Route pour recevoir le fichier image collé et faire l'OCR + matching."""
    if 'file' not in request.files:
        return jsonify({"error": "Aucun fichier envoyé."}), 400
    file = request.files['file']
    filename = file.filename if file.filename else "pasted_image.png"
    path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(path)

    product_names = extract_product_names_from_image(path)
    matched_parfums = match_names_to_json(product_names)
    return jsonify({"produits_detectes": [p['nom'] for p in matched_parfums]})

@app.route("/generate_brief", methods=["POST"])
def generate_brief():
    """Route pour générer le document Word final."""
    data = request.json
    periode = data['periode']
    type_doc = data['typeDoc']
    parfums = data['parfums']  # liste de noms (ex: ["Etoile", "Concerto"])

    # Choisir le template
    template_path = choose_template(type_doc)

    # Appeler la fonction qui remplit le template
    output_path = fill_template(template_path, periode, parfums=[])

    # Note : si tu veux reparser 'parfums' pour en tirer plus de data, fais-le ici

    # Renvoyer le fichier
    return send_file(output_path, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True, port=5000)
